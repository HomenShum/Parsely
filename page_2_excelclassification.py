def excelclassification_tool():
    import asyncio
    import logging
    import os
    import time
    from asyncio import Semaphore
    from io import BytesIO
    from typing import Any, Dict, List, Optional

    import pandas as pd
    import streamlit as st
    import streamlit_antd_components as sac
    from icecream import ic
    from openai import AsyncOpenAI
    from openpyxl import load_workbook
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter
    from pydantic import BaseModel
    from tenacity import retry, stop_after_attempt, wait_fixed, retry_if_exception_type
    from openai import RateLimitError

    import openai
    import instructor
    from pydantic import BaseModel
    from pydantic import ValidationError

    # Configure logging
    logging.basicConfig(level=logging.DEBUG)

    # Patch the OpenAI client with Instructor, 062824 Instructor Version Update
    # aclient = instructor.apatch(AsyncOpenAI(api_key = st.secrets["OPENAI_API_KEY"]))
    aclient = instructor.from_openai(openai.AsyncOpenAI(api_key = st.secrets["OPENAI_API_KEY"]))

    # Define Pydantic model for response validation
    class CompanySectorTagClass(BaseModel):
        SectorTag: str
        QuickInfo: List[str]

    class CompanyClassificationBool(BaseModel):
        Result: bool
        Reason: str

    class CompanyClassificationGeneral(BaseModel):
        Response: List[str]

    class CompanyClassificationGeneral2(BaseModel):
        Response: str

    # Asynchronous function to query data using OpenAI and validate with Pydantic
    sem = asyncio.Semaphore(1000)
    retry_decorator = retry(
        stop=stop_after_attempt(3),
        wait=wait_fixed(2),
        retry=retry_if_exception_type(RateLimitError),
        reraise=True
    )

    @retry_decorator
    async def rate_limited_company_classification_async(company_data: Dict[str, Any], sem: Semaphore, prompt: str) -> List[str]:
        # Determine the response model based on classification choice
        if st.session_state['classification_choice'] == 'sector-tag':
            response_model = CompanySectorTagClass
        elif st.session_state['classification_choice'] == 'yes-no-reasoning':
            response_model = CompanyClassificationBool
        elif st.session_state['classification_choice'] == 'general_response':
            response_model = CompanyClassificationGeneral

        model_choice = st.session_state['model_choice']

        try:
            async with sem:
                try:
                    model = await aclient.chat.completions.create(
                        model=model_choice,
                        response_model=response_model,
                        messages=[
                            {"role": "user", "content": " System Prompt: " + str(company_data) + " User Prompt: " + prompt},
                        ],
                        max_retries=3,
                    )
                    # Process the response based on classification choice
                    if st.session_state['classification_choice'] == 'sector-tag':
                        return [model.SectorTag, model.QuickInfo]
                    elif st.session_state['classification_choice'] == 'yes-no-reasoning':
                        return [model.Result, model.Reason]
                    elif st.session_state['classification_choice'] == 'general_response':
                        return [model.Response]
                except ValidationError as ve:
                    logging.error(f"Pydantic validation error: {ve}")
                    st.error(f"Pydantic validation error: {ve}")
                    raise ve
                except Exception as e:
                    logging.error(f"Error during classification: {e}")
                    st.error(f"Error during classification: {e}")
                    raise
        except RateLimitError as e:
            logging.error(f"Rate limit exceeded: {e}")
            st.warning("Rate limit exceeded. Retrying after a delay...")
            await asyncio.sleep(60)  # Wait before retrying
            raise e
        except Exception as e:
            logging.error(f"Error during classification: {e}")
            st.error(f"An error occurred during classification: {e}")
            raise e


    async def main_classification_async(df, prompt):
        if 'PII Masked Data' not in df.columns:
            # Concatenate all columns into a single string
            df_dict_list = df.apply(lambda x: ', '.join([f'{column}: {x[column]}' for column in df.columns]), axis=1).to_list()
        else:
            df_dict_list = df['PII Masked Data'].to_dict('records')

        tasks = [rate_limited_company_classification_async(data, sem, prompt) for data in df_dict_list]
        total_tasks = len(tasks)
        progress_bar = st.progress(0)
        results = []

        for idx, task in enumerate(asyncio.as_completed(tasks)):
            try:
                result = await task
                results.append(result)
            except Exception as e:
                logging.error(f"Task resulted in an exception: {e}")
                results.append(None)
            progress_bar.progress((idx + 1) / total_tasks)

        # Add new columns to the dataframe df
        if st.session_state['classification_choice'] == 'sector-tag':
            results_df = pd.DataFrame(results, columns=['Sector Tag', 'QuickInfo'])
        elif st.session_state['classification_choice'] == 'yes-no-reasoning':
            results_df = pd.DataFrame(results, columns=['Result', 'Reason'])
        elif st.session_state['classification_choice'] == 'general_response':
            results_df = pd.DataFrame(results, columns=['Response'])
        df = pd.concat([st.session_state['new_df_to_be_processed'], results_df], axis=1)


        # Cleanup process starts here
        results_df_dict_list = results_df.to_dict('records')
        clean_up_prompt_list = []
        for i, result_dict in enumerate(results_df_dict_list):
            clean_up_prompt_list.append((i, f"{prompt}. Result data: {result_dict}"))

        clean_up_tasks = [(i, rate_limited_clean_up_async(data, sem, prompt)) for i, data in clean_up_prompt_list]
        clean_up_results = await asyncio.gather(*[task for _, task in clean_up_tasks])

        clean_up_results_df = pd.DataFrame(clean_up_results, columns=['Cleaned Up Result'])
        clean_up_final_df = pd.concat([df, clean_up_results_df], axis=1)

        return clean_up_final_df


    @retry_decorator
    async def rate_limited_query_main_generate_new_col_async(data: Dict[str, Any], sem: Semaphore, prompt: str) -> List[str]:
        if 'model_choice' not in st.session_state:
            st.session_state['model_choice'] = "gpt-4o-mini"  # Initialize with a default value

        model_choice = st.session_state['model_choice']
        async with sem:
            # print(f"Processing description: {str(data)}")
            model = await aclient.chat.completions.create(
                # model="gpt-4o-mini",
                model=model_choice,
                response_model=CompanyClassificationGeneral2,
                messages=[
                    {"role": "user", "content": + " System Prompt: " + prompt_for_generating_new_column + " User Prompt: " + str(data)},
                ],
            max_retries=3,
            )
            # ic(model)
            # return model.choices[0].message.content
            return [model.Response]

    @retry_decorator
    async def rate_limited_clean_up_async(data: dict, sem: asyncio.Semaphore, prompt: str) -> str:
        async with sem:
            model = await aclient.chat.completions.create(
                model="gpt-4o-mini",
                response_model=CompanyClassificationGeneral2,
                messages=[
                    {"role": "user", "content": "Clean up the irrelevant result. Keep only useful and relevant information in readable format. Double Check and Proof Read." + " User Prompt: " + str(data)},
                ],
                max_retries=3,
            )
            return [model.Response]

    async def main_generate_new_col_async(df_a, df_b, prompt):
        df_a_dict_list = df_a.to_dict('records')
        df_b_dict_list = df_b.to_dict('records')

        prompt_list = []
        batch_size = st.session_state.batch_size  # Adjust this value based on your needs
        for i in range(len(df_b_dict_list)):
            for j in range(0, len(df_a_dict_list), batch_size):  # Generate results for each batch in df_a_dict_list
                df_a_dict_batch = df_a_dict_list[j:j+batch_size]  # Select a batch from df_a_dict_list in order
                prompt_list.append((i, f"Sheet A data: {df_a_dict_batch}. Sheet B data: {df_b_dict_list[i]}"))

        tasks = [(i, rate_limited_query_main_generate_new_col_async(data, sem, prompt)) for i, data in prompt_list]
        results = await asyncio.gather(*[task for _, task in tasks])

        # Group the results by the index of df_b_dict_list
        results_grouped = [[] for _ in range(len(df_b_dict_list))]
        for (i, _), result in zip(tasks, results):
            results_grouped[i].extend(result)

        # Flatten the results_grouped list and convert it into DataFrame
        results_grouped_flat = [" ".join(map(str, result)) for result in results_grouped]
        results_df = pd.DataFrame(results_grouped_flat, columns=['Generated New Column'])

        final_df = pd.concat([st.session_state['new_df_to_be_processed_2'], results_df], axis=1)

        # Clean up the results_df that has the combined results
        results_df_dict_list = results_df.to_dict('records')
        clean_up_prompt_list = []
        for i in range(len(results_df_dict_list)):
            clean_up_prompt_list.append((i, f"{prompt}. Sheet B data: {df_b_dict_list[i]}. Final result: {results_df_dict_list[i]}"))

        clean_up_tasks = [(i, rate_limited_clean_up_async(data, sem, prompt)) for i, data in clean_up_prompt_list]
        clean_up_results = await asyncio.gather(*[task for _, task in clean_up_tasks])

        clean_up_results_df = pd.DataFrame(clean_up_results, columns=['Cleaned Up Result'])
        
        clean_up_final_df = pd.concat([final_df, clean_up_results_df], axis=1)

        return clean_up_final_df
        

    def unmask_data(masked_data, unmasking_dict):
        unmasked_data = masked_data
        for masked, unmasked in unmasking_dict.items():
            unmasked_data = unmasked_data.replace(masked, unmasked)
        return unmasked_data

########################################################################################
    # st.set_page_config(layout="wide")

    from docx import Document
    import pandas as pd
    import csv
    from io import StringIO, BytesIO
    import os
    from bs4 import BeautifulSoup
    import re
    # from unstructured.partition.auto import partition
    import tempfile
    import aiohttp
    import asyncio
    import aiofiles
    import logging

    @st.cache_data
    def is_html(text):
        return bool(re.search(r'<[a-zA-Z][\s\S]*>', text))

    import aiohttp
    import aiofiles
    import os
    import logging
    import re
    import streamlit as st



    # def read_docx_pages(docx_file):
    #     # Save the UploadedFile to a temporary file
    #     with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
    #         # Assuming docx_file is a file-like object; if not, adjust accordingly
    #         tmp.write(docx_file.read())
    #         tmp_path = tmp.name

    #     # Use the temporary file path with Document
    #     doc = Document(tmp_path)
    #     full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
    #     if is_html(full_text):
    #         pages = []
    #         for paragraph in doc.paragraphs:
    #             if paragraph.text.strip():  # Check if paragraph is not just whitespace
    #                 soup = BeautifulSoup(paragraph.text, 'html.parser')
    #                 cleaned_text = soup.get_text().strip()
    #                 if cleaned_text:  # Filter out empty rows
    #                     pages.append(cleaned_text)
    #     else:
    #         # Now, use the path of the temporary file with partition
    #         elements = partition(tmp_path)
    #         pages = []

    #         for element in elements:
    #             text = str(element).strip()
    #             if text:  # Filter out empty rows
    #                 pages.append(text)

    #     # Optionally, delete the temporary file if no longer needed
    #     os.remove(tmp_path)

    #     return pages


    # Streamlit UI for uploading and downloading

    st.title("DOCX to CSV/Parquet/Excel Page Extractor")

    # Define fragments for partial rerun
    @st.experimental_fragment()
    def upload_docx_files():
        return st.file_uploader("Upload DOCX files", type=["docx"], accept_multiple_files=True)

    @st.experimental_fragment()
    @st.cache_data
    def display_conversion_options(uploaded_docx_files):
        async def aupload_file_to_api(file_path):
            url = 'https://txtparseapis-azure.ashysky-c2a561fc.westus2.azurecontainerapps.io/process_upload'
            headers = {'accept': 'application/json'}
            # Ensure the file exists and is a file
            if os.path.isfile(file_path):
                async with aiohttp.ClientSession() as session:
                    # Open the file in binary mode
                    async with aiofiles.open(file_path, 'rb') as f:
                        file_content = await f.read()
                    data = aiohttp.FormData()
                    data.add_field('file', file_content, filename=os.path.basename(file_path), content_type='text/plain')
                    async with session.post(url, data=data, headers=headers) as response:
                        if response.status == 200:
                            response_json = await response.json()
                            # Ensure the return value is serializable
                            return dict(response_json)  # Convert to dict if not already, ensuring serializability
                        else:
                            logging.error(f"Failed to upload file. Status: {response.status}")
                            return {"error": f"Failed to upload file. Status: {response.status}"}
            else:
                logging.error("File does not exist or is not a file.")
                return {"error": "File does not exist or is not a file."}

        async def aread_docx_pages(docx_files):
            tasks = [aprocess_single_file(docx_file) for docx_file in docx_files]
            pages_list = await asyncio.gather(*tasks)
            return pages_list

        async def aprocess_single_file(docx_file):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(docx_file.read())
                tmp_path = tmp.name

            doc = Document(tmp_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

            if is_html(full_text):
                pages = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        soup = BeautifulSoup(paragraph.text, 'html.parser')
                        cleaned_text = soup.get_text().strip()
                        if cleaned_text:
                            pages.append(cleaned_text)
            else:
                response_json = await aupload_file_to_api(tmp_path)
                if response_json:
                    # Assuming the API returns a list of pages/text elements
                    pages = response_json.get('pages', [])
                else:
                    pages = []

            os.remove(tmp_path)
            return pages

        # Function to save pages to Excel
        @st.cache_data
        def save_to_excel(pages):
            df = pd.DataFrame({'extracted_content': pages})
            output = BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False, sheet_name='Content')
            return output.getvalue()

        # Function to save pages to CSV, all content in one column
        @st.cache_data
        def save_to_csv(pages):
            output = StringIO()
            writer = csv.writer(output, quoting=csv.QUOTE_ALL)  # Use QUOTE_ALL to enclose content in quotes
            writer.writerow(["extracted_content"])
            for page in pages:
                writer.writerow([page])
            return output.getvalue()

        # Function to save pages to Parquet
        @st.cache_data
        def save_to_parquet(pages):
            df = pd.DataFrame({'extracted_content': pages})
            output = BytesIO()
            df.to_parquet(output, index=False)
            return output.getvalue()

        if uploaded_docx_files:
            st.write("Files uploaded successfully.")
            pages_list = asyncio.run(aread_docx_pages(uploaded_docx_files))
            
            # Calculate the number of rows needed for the grid
            num_files = len(uploaded_docx_files)
            num_rows = (num_files + 2) // 3  # Ensure there's a row for any remainder
            
            for row in range(num_rows):
                cols = st.columns(3)  # Create a new row of columns
                for i in range(3):
                    file_index = row * 3 + i
                    if file_index < num_files:
                        uploaded_docx_file = uploaded_docx_files[file_index]
                        pages = pages_list[file_index]
                        with cols[i]:
                            with st.expander(f"{uploaded_docx_file.name}"):
                                st.subheader(f"{uploaded_docx_file.name}")
                                csv_content = save_to_csv(pages)
                                original_filename = os.path.splitext(uploaded_docx_file.name)[0]
                                csv_filename = f"{original_filename}.csv"
                                st.download_button("Download CSV", csv_content, file_name=csv_filename, mime="text/csv")
                                parquet_content = save_to_parquet(pages)
                                parquet_filename = f"{original_filename}.parquet"
                                st.download_button("Download Parquet", parquet_content, file_name=parquet_filename, mime="application/octet-stream")
                                excel_content = save_to_excel(pages)
                                excel_filename = f"{original_filename}.xlsx"
                                st.download_button("Download Excel", excel_content, file_name=excel_filename, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    uploaded_docx_files = upload_docx_files()
    st.subheader("Converted Files and Download Links")
    display_conversion_options(uploaded_docx_files)

    # if uploaded_docx_file:
    #     st.write("File uploaded successfully.")
    #     pages = read_docx_pages(uploaded_docx_file)
        
    #     # Convert to CSV
    #     csv_content = save_to_csv(pages)
    #     original_filename = os.path.splitext(uploaded_docx_file.name)[0]
    #     csv_filename = f"{original_filename}.csv"
    #     st.download_button("Download CSV", csv_content, file_name=csv_filename, mime="text/csv")
        
    #     # Convert to Parquet
    #     parquet_content = save_to_parquet(pages)
    #     parquet_filename = f"{original_filename}.parquet"
    #     st.download_button("Download Parquet", parquet_content, file_name=parquet_filename, mime="application/octet-stream")
        
    #     # Convert to Excel
    #     excel_content = save_to_excel(pages)
    #     excel_filename = f"{original_filename}.xlsx"
    #     st.download_button("Download Excel", excel_content, file_name=excel_filename, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.divider()

    st.title("Excel/CSV/Parquet Classification Tool")

    uploaded_xlsx_csv_parquet_files = st.file_uploader("Upload an Excel, CSV, or Parquet file", type=["csv", "xlsx", "parquet", "xlsm"], accept_multiple_files=True, key="unique_file_uploader_key")
    
    if uploaded_xlsx_csv_parquet_files:
        for uploaded_file in uploaded_xlsx_csv_parquet_files:
            file_extension = os.path.splitext(uploaded_file.name)[1]
            
            if file_extension == '.csv':
                df = pd.read_csv(uploaded_file)
            elif file_extension in ['.xlsx', '.xlsm']: 
                df = pd.read_excel(uploaded_file)
            elif file_extension == '.parquet':
                df = pd.read_parquet(uploaded_file)
            
            st.write(f"Displaying content of {uploaded_file.name}:")
            st.dataframe(df)

    # st.divider()

    # st.title("Excel/CSV Classification Tool")

    # uploaded_xlsx_csv_files = st.file_uploader("Upload an Excel or CSV file", type=["csv", "xlsx"], accept_multiple_files=True)

    st.divider()

    selected_modes = sac.buttons([
        sac.ButtonsItem(label='Download'),
        sac.ButtonsItem(label='Classification'),
        sac.ButtonsItem(label='Build Your Own Bot'),
        sac.ButtonsItem(label='Generate New Column on Sheet B Given A')
    ], label='Select Use Cases', align='start')

    col1, col2 = st.columns([1, 2])

    if 'new_df_to_be_processed' not in st.session_state:
        st.session_state['new_df_to_be_processed'] = pd.DataFrame()

    if 'new_df_to_be_processed_2' not in st.session_state:
        st.session_state['new_df_to_be_processed_2'] = pd.DataFrame()


    # Check if a file has been uploaded
    if uploaded_xlsx_csv_parquet_files and ('files_processed' not in st.session_state or len(uploaded_xlsx_csv_parquet_files) != len(st.session_state['uploaded_dfs'])):
        st.session_state['uploaded_dfs'] = []

        for user_file in uploaded_xlsx_csv_parquet_files:
            try:
                if user_file.name.endswith(('xlsx', 'xlsm')):
                    excel_file = pd.ExcelFile(user_file)
                    sheet_names = excel_file.sheet_names
                    sheet_data = []

                    for sheet in sheet_names:
                        df = pd.read_excel(excel_file, sheet_name=sheet)
                        rows = df.to_dict('index')
                        sheet_data.append({
                            'SheetName': sheet,
                            'Rows': rows,
                            'DataFrame': df
                        })

                    st.session_state['uploaded_dfs'].append({
                        'FileName': user_file.name,
                        'Sheets': sheet_data
                    })
                    
                elif user_file.name.endswith('csv'):
                    df = pd.read_csv(user_file)
                    rows = df.to_dict('index')
                    sheet_data = [{'SheetName': 'na', 'Rows': rows, 'DataFrame': df}]

                    st.session_state['uploaded_dfs'].append({
                        'FileName': user_file.name,
                        'Sheets': sheet_data
                    })

                elif user_file.name.endswith('parquet'):
                    df = pd.read_parquet(user_file)
                    rows = df.to_dict('index')
                    sheet_data = [{'SheetName': 'na', 'Rows': rows, 'DataFrame': df}]

                    st.session_state['uploaded_dfs'].append({
                        'FileName': user_file.name,
                        'Sheets': sheet_data
                    })

                else:
                    st.error('Unsupported file type. Please upload a .csv, .xlsx, .xlsm, or .parquet file.')

            except Exception as e:
                st.error(f'Error reading file: {e}')

        with col1:
            # Allow the user to select a file and a sheet
            if 'uploaded_dfs' in st.session_state:
                file_names = [file['FileName'] for file in st.session_state['uploaded_dfs']]

                # If 'selected_file' is not in the session state, initialize it to the first file name
                if 'selected_file' not in st.session_state:
                    if file_names:  # Check if file_names is not empty
                        st.session_state['selected_file'] = file_names[0]
                    else:
                        st.session_state['selected_file'] = None

                st.divider()

                st.header('Select a file as sheet A')
                
                selected_file = st.selectbox('Select a file', file_names, key='selected_file')

                if selected_file:
                    selected_file_data = next((file for file in st.session_state['uploaded_dfs'] if file["FileName"] == selected_file), None)
                    if selected_file_data:
                        sheet_names = [sheet['SheetName'] for sheet in selected_file_data['Sheets']]

                        # If the selected file has changed, reset the selected sheet
                        if st.session_state['selected_file'] != selected_file:
                            st.session_state['selected_sheet'] = sheet_names[0]

                        # If 'selected_sheet' is not in the session state, initialize it to the first sheet name
                        if 'selected_sheet' not in st.session_state:
                            st.session_state['selected_sheet'] = sheet_names[0]

                        selected_sheet = st.selectbox('Select a sheet', sheet_names, key='selected_sheet')
                        df = next((sheet['DataFrame'] for sheet in selected_file_data['Sheets'] if sheet["SheetName"] == selected_sheet), None)

                        # Use st.data_editor to display the DataFrame
                        if df is not None:
                            st.write('Data Preview:')
                            st.dataframe(df)

                            # Allow the user to select the column names
                            column_names = df.columns.tolist()
                            chip_items = [sac.ChipItem(label=col) for col in column_names]

                            # Add a checkbox for the "Select All" option
                            select_all = st.checkbox('Select All')
                            
                            # Initialize new_df as an empty DataFrame
                            new_df = pd.DataFrame()

                            if select_all:
                                selected_columns = column_names
                                st.session_state['new_df_to_be_processed'] = df
                            else:
                                selected_columns = sac.chip(
                                    items=chip_items, 
                                    label='Select the columns to be downloaded or processed:', 
                                    align='start',
                                    radius='md', 
                                    multiple=True
                                )

                                if not select_all and not selected_columns:
                                    selected_columns = []
                                elif selected_columns:
                                    new_df = df[selected_columns]

                                st.session_state['new_df_to_be_processed'] = new_df

                            # If the user clicked a button, output a new Excel file
                            output = BytesIO()
                            new_df.to_excel(output, index=False)

                            # Load the workbook
                            wb = load_workbook(output)

                            # Select the active worksheet
                            ws = wb.active

                            # Format the cells
                            for row in ws.iter_rows():
                                for cell in row:
                                    cell.font = Font(bold=True)

                            # Adjust the width of the columns
                            max_column_width = 100  # Set the maximum column width
                            for column in ws.columns:
                                max_length = 0
                                column = [cell for cell in column]
                                for cell in column:
                                    try:
                                        if len(str(cell.value)) > max_length:
                                            max_length = len(cell.value)
                                    except:
                                        pass
                                adjusted_width = min((max_length + 2), max_column_width)
                                ws.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width

                            # Save the workbook
                            output = BytesIO()
                            wb.save(output)

                            st.session_state['excel_output'] = output.getvalue()
                            st.session_state['selected_file_name'] = f"{selected_file.split('.')[0]}_parsely_output.xlsx"
            if not st.session_state['new_df_to_be_processed'].empty:
                st.write(st.session_state['new_df_to_be_processed'])

                st.divider()

                if 'Download' in selected_modes:
                    excel_download_button = st.download_button(
                        label="Download Excel",
                        data=st.session_state['excel_output'],
                        file_name=st.session_state['selected_file_name'],
                        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    )

                    if excel_download_button:
                        st.success('File downloaded successfully')


                # rather than using the entire new_df, user may choose which data to use for classification
                st.subheader("Select the columns used by the LLM (Less columns; Less word token costs):")
                select_all_for_classification = st.checkbox('Select All', key='select_all_for_classification')

                column_names = st.session_state['new_df_to_be_processed'].columns.tolist()
                chip_items = [sac.ChipItem(label=col) for col in column_names]

                new_df_for_classification = pd.DataFrame()
                
                if select_all_for_classification:
                    st.session_state['new_df_to_be_processed_by_llm'] = st.session_state['new_df_to_be_processed']
                else:
                    selected_columns_for_classification = sac.chip(
                        items=chip_items, 
                        label='Select the columns to be used:', 
                        align='start',
                        radius='md', 
                        multiple=True
                    )

                    if not select_all_for_classification and not selected_columns_for_classification:
                        selected_columns_for_classification = []
                    elif selected_columns_for_classification:
                        new_df_for_classification = st.session_state['new_df_to_be_processed'][selected_columns_for_classification]

                    st.session_state['new_df_to_be_processed_by_llm'] = new_df_for_classification

                if not st.session_state['new_df_to_be_processed_by_llm'].empty:
                    st.markdown(f"**Using the following columns:** \n\n{st.session_state['new_df_to_be_processed_by_llm'].columns.tolist()}")
                    df_dict_list = st.session_state['new_df_to_be_processed_by_llm'].to_dict('records')
                    st.markdown(f"**Preview header row used:**\n\n'{df_dict_list[0]}'")


                if 'Classification' in selected_modes:
                    classification_prompt = """
                        1. Classify the company into a sector tag based on the description.
                            Select from one of the five tags: 
                                BioPharma: Companies that develop drugs, vaccines, and other therapeutics for human use. 
                                MedTech: Companies that develop medical devices, including diagnostics, imaging, and surgical tools. 
                                HealthTech: Companies that develop software and other healthcare technologies to improve healthcare delivery. 
                                Diagnostics: Companies that develop diagnostic tools, including imaging, genetic testing, and other assays. 
                                MedicalTools: Companies that develop tools for medical research, including lab equipment and reagents. 
                            Other: If the company does not fit into any of the above categories, select this option.
                        2. Quick info is a short list of keywords that summarizes the company purpose based on the description. Schema: Keyword1, Keyword2, Keyword3, Keyword4, Keyword5
                    """

                    st.subheader("Edit the prompt for classification:")

                    edited_classification_prompt = st.text_area(
                        label = "Please edit the prompt if necessary and press Ctrl+Enter to continue:", 
                        value = classification_prompt,
                        height = 250
                    )

                    if edited_classification_prompt:
                        st.session_state['edited_classification_prompt'] = edited_classification_prompt
                        st.write("Here is your 'Prompt' used for classification.")
                        st.write(st.session_state['edited_classification_prompt'])


                    # if st.button("Start PII Masking!"):
                    #     # Mask PII in the data, look through each row  
                    #     start_time = time.time()
                    #     post_piim_df = pd.DataFrame(columns=['PII Masked Data', 'Unmasking Dictionary', 'Unmasked Data'])
                    #     for index, row in new_df.iterrows():
                    #         # combine all row information from every columns. the result should be "column1:row1, column2:row1, column3:row1, ..."
                    #         # pre_piim_text = ', '.join([f"{column}:{row[column]}" for column in new_df.columns])
                    #         pre_piim_text = ', '.join([f"{column}:{int(row[column]) if pd.to_numeric(row[column], errors='coerce') % 1 == 0 else row[column]}" for column in new_df.columns])
                    #         node = TextNode(text=pre_piim_text)
                    #         processor = NERPIINodePostprocessor()
                    #         new_nodes = processor.postprocess_nodes([NodeWithScore(node=node)])
                    #         masked_data = new_nodes[0].node.get_text()
                    #         unmasking_dict = new_nodes[0].node.metadata["__pii_node_info__"]
                    #         unmasked_data = unmask_data(masked_data, unmasking_dict)
                    #         new_row = pd.DataFrame({'PII Masked Data': [masked_data], 'Unmasking Dictionary': [unmasking_dict], 'Unmasked Data': [unmasked_data]})
                    #         post_piim_df = pd.concat([post_piim_df, new_row], ignore_index=True)
                    #     # Preview the masked data
                    #     st.write(post_piim_df)
                    #     st.session_state['post_piim_df'] = post_piim_df
                    #     st.success('PII masking completed successfully')
                    #     st.success(f"Processing time: {time.time() - start_time} seconds.")
                    if st.button("Start PII Masking!"):
                        # Mask PII in the data, look through each row  
                        start_time = time.time()
                        post_piim_df = pd.DataFrame(columns=['PII Masked Data', 'Unmasking Dictionary', 'Unmasked Data'])
                        for index, row in st.session_state['new_df_to_be_processed'].iterrows():
                            masked_row = []
                            unmasking_dict = {}
                            for column in st.session_state['new_df_to_be_processed'].columns:
                                # create a unique identifier for each data point based on its column name and row index
                                pre_piim_text = f"{column}_{index}"
                                masked_data = f"MASKED_{pre_piim_text}"
                                masked_row.append(masked_data)
                                unmasking_dict[masked_data] = row[column]
                            # concatenate all the masked data in the row into a single string
                            masked_row_str = ', '.join(masked_row)
                            unmasked_data = ', '.join([str(row[column]) for column in st.session_state['new_df_to_be_processed'].columns])
                            new_row = pd.DataFrame({'PII Masked Data': [masked_row_str], 'Unmasking Dictionary': [unmasking_dict], 'Unmasked Data': [unmasked_data]})
                            post_piim_df = pd.concat([post_piim_df, new_row], ignore_index=True)
                        # Preview the masked data
                        st.write(post_piim_df)
                        st.session_state['post_piim_df'] = post_piim_df
                        st.session_state['did_piim'] = True
                        st.success('PII masking completed successfully')
                        st.success(f"Processing time: {time.time() - start_time} seconds.")
                    else:
                        st.warning('If you have not masked the PII data, then the selected data or all data will be used for classification.')
                        st.session_state['did_piim'] = False        

                    if 'general_classification' not in st.session_state:
                        st.session_state.general_classification = False

                    # Choose between general classification and sector classification
                    st.subheader("Select the classification type:")
                    classification_choice = st.selectbox('Select the classification type', ['sector-tag', 'yes-no-reasoning', 'general_response'], key='classification_choice')

                    # Select model: gpt-4-turbo, gpt-4o-mini
                    st.subheader("Select the model:")
                    model_choice = st.selectbox('Select the model', ['gpt-4o-mini', 'gpt-4o-mini'], key='model_choice')
                    # print(model_choice)

                    if st.button("Looks Good! Start Classification"):
                        # Start the classification process
                        start_time = time.time()
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        if st.session_state['did_piim']:
                            results_df = loop.run_until_complete(main_classification_async(st.session_state['post_piim_df'], st.session_state['edited_classification_prompt']))
                        else:
                            results_df = loop.run_until_complete(main_classification_async(st.session_state['new_df_to_be_processed_by_llm'], st.session_state['edited_classification_prompt']))
                        st.write(f"Processing time: {time.time() - start_time} seconds.")
                        
                        # Display the results DataFrame
                        st.dataframe(results_df)

                        # store in session_state
                        st.session_state['classified_results_df'] = results_df

                    # if session state not empty, download the results
                    if 'classified_results_df' in st.session_state:
                        classified_results_df = st.session_state['classified_results_df']

                        # If the user clicked a button, output a new Excel file
                        output = BytesIO()
                        classified_results_df.to_excel(output, index=False)

                        # Load the workbook
                        wb = load_workbook(output)
                        
                        # Select the active worksheet
                        ws = wb.active

                        # Format the cells
                        for row in ws.iter_rows():
                            for cell in row:
                                cell.font = Font(bold=True)

                        # Adjust the width of the columns
                        max_column_width = 100  # Set the maximum column width
                        for column in ws.columns:
                            max_length = 0
                            column = [cell for cell in column]
                            for cell in column:
                                try:
                                    if len(str(cell.value)) > max_length:
                                        max_length = len(cell.value)
                                except:
                                    pass
                            adjusted_width = min((max_length + 2), max_column_width)
                            ws.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width


                        # Save the workbook
                        output = BytesIO()
                        wb.save(output)

                        st.session_state['excel_output'] = output.getvalue()
                        st.session_state['selected_file_name'] = f"{selected_file.split('.')[0]}_parsely_classification_output.xlsx"

                        # Create the download button
                        excel_download_button = st.download_button(
                            label="Download Classification Results",
                            data=st.session_state['excel_output'],
                            file_name=st.session_state['selected_file_name'],
                            mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        )

                        if excel_download_button:
                            st.success('File downloaded successfully')

        if 'Build Your Own Bot' in selected_modes:
            st.write("Building Your Own Bot with DataFrame")

        if 'selected_file_2' not in st.session_state:
            st.session_state['selected_file_2'] = None
        if 'selected_sheet_2' not in st.session_state:
            st.session_state['selected_sheet_2'] = None
        if 'new_df_to_be_processed_2' not in st.session_state:
            st.session_state['new_df_to_be_processed_2'] = pd.DataFrame()
        if 'new_df_to_be_processed_by_llm_2' not in st.session_state:
            st.session_state['new_df_to_be_processed_by_llm_2'] = pd.DataFrame()
        if 'batch_size' not in st.session_state:
            st.session_state['batch_size'] = 5


        if 'Generate New Column on Sheet B Given A' in selected_modes:
            with col2:
                st.divider()
                st.header('Select a second file as sheet B')
                
                selected_file_2 = st.selectbox('Select a file', file_names, key='selected_file_2')
                selected_file_data_2 = next((file for file in st.session_state['uploaded_dfs'] if file["FileName"] == selected_file_2), None)
                if selected_file_data_2:
                    sheet_names_2 = [sheet['SheetName'] for sheet in selected_file_data_2['Sheets']]

                    if st.session_state['selected_file_2'] != selected_file_2:
                        st.session_state['selected_sheet_2'] = sheet_names_2[0]

                    if 'selected_sheet_2' not in st.session_state:
                        st.session_state['selected_sheet_2'] = sheet_names_2[0]
                        
                    selected_sheet_2 = st.selectbox('Select a sheet', sheet_names_2, key='selected_sheet_2')
                    df_2 = next((sheet['DataFrame'] for sheet in selected_file_data_2['Sheets'] if sheet["SheetName"] == selected_sheet_2), None)

                    if df_2 is not None:
                        st.write('Data Preview:')
                        st.dataframe(df_2)

                        # Allow the user to select the column names
                        column_names_2 = df_2.columns.tolist()
                        chip_items_2 = [sac.ChipItem(label=col) for col in column_names_2]

                        # Add a checkbox for the "Select All" option
                        select_all_2 = st.checkbox('Select All', key='select_all_2')

                        # Initialize new_df as an empty DataFrame
                        new_df_2 = pd.DataFrame()

                        if select_all_2:
                            selected_columns_2 = column_names_2
                            st.session_state['new_df_to_be_processed_2'] = df_2
                        else:
                            selected_columns_2 = sac.chip(
                                items=chip_items_2, 
                                label='Select the 2nd columns to be downloaded or processed:', 
                                align='start',
                                radius='md', 
                                multiple=True,
                                key='selected_columns_2_chip'  # Unique key for this widget
                            )

                            if not select_all_2 and not selected_columns_2:
                                selected_columns_2 = []
                            elif selected_columns_2:
                                new_df_2 = df_2[selected_columns_2]

                            st.session_state['new_df_to_be_processed_2'] = new_df_2

                if not st.session_state['new_df_to_be_processed_2'].empty:
                    st.write(st.session_state['new_df_to_be_processed_2'])

                    st.divider()

                    st.subheader("Select the columns used by the LLM (Less columns; Less word token costs):")
                    select_all_for_classification_2 = st.checkbox('Select All', key='select_all_for_classification_2')

                    column_names_2 = st.session_state['new_df_to_be_processed_2'].columns.tolist()
                    chip_items_2 = [sac.ChipItem(label=col) for col in column_names_2]

                    new_df_for_classification_2 = pd.DataFrame()

                    if select_all_for_classification_2:
                        st.session_state['new_df_to_be_processed_by_llm_2'] = st.session_state['new_df_to_be_processed_2']
                    else:
                        selected_columns_for_classification_2 = sac.chip(
                            items=chip_items_2, 
                            label='Select the columns to be used:', 
                            align='start',
                            radius='md', 
                            multiple=True,
                            key='selected_columns_for_classification_2_chip'  # Unique key for this widget
                        )

                        if not select_all_for_classification_2 and not selected_columns_for_classification_2:
                            selected_columns_for_classification_2 = []
                        elif selected_columns_for_classification_2:
                            new_df_for_classification_2 = st.session_state['new_df_to_be_processed_2'][selected_columns_for_classification_2]

                        st.session_state['new_df_to_be_processed_by_llm_2'] = new_df_for_classification_2

                    if not st.session_state['new_df_to_be_processed_by_llm_2'].empty:
                        st.markdown(f"**Using the following columns:** \n\n{st.session_state['new_df_to_be_processed_by_llm_2'].columns.tolist()}")
                        df_dict_list_2 = st.session_state['new_df_to_be_processed_by_llm_2'].to_dict('records')
                        st.markdown(f"**Preview header row used :**\n\n'{df_dict_list_2[0]}'")

            st.divider()

            st.header('Configure prompt for generating new column on Sheet B given A')

            prompt_for_generating_new_column = st.text_area(
                label = "Please edit the prompt if necessary and press Ctrl+Enter to continue:", 
                value = "Extract contact names and emails that matches Sheet B's company name. If there is no match, then leave blank.",
                height = 100
            )

            # Display a preview of the prompt composition
            if st.session_state['new_df_to_be_processed_by_llm'].empty or st.session_state['new_df_to_be_processed_by_llm_2'].empty:
                st.warning("Please select columns for both sheets before generating the prompt preview.")
            else:
                df_a_dict_list = st.session_state['new_df_to_be_processed_by_llm'].to_dict('records')
                df_b_dict_list = st.session_state['new_df_to_be_processed_by_llm_2'].to_dict('records')
                batch_size = st.session_state.batch_size
                i = 0
                j = 0
                df_a_dict_batch = df_a_dict_list[j:j+batch_size]
                prompt_preview = f"Sheet A data: {df_a_dict_batch}. Sheet B data: {df_b_dict_list[i]}"
                st.markdown(f"**Prompt Preview:**\n\n{prompt_preview}")

            if st.button("Looks Good! Start Generating New Column on Sheet B"):
                start_time = time.time()
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    new_column_results_df = loop.run_until_complete(main_generate_new_col_async(st.session_state['new_df_to_be_processed_by_llm'], st.session_state['new_df_to_be_processed_by_llm_2'], prompt_for_generating_new_column))
                except RateLimitError:
                    st.info("Rate limit reached for the OpenAI API. Please try again later.")
                st.write(f"Processing time: {time.time() - start_time} seconds.")
                
                st.dataframe(new_column_results_df)

                st.session_state['new_column_results_df'] = new_column_results_df
